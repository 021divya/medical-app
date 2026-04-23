import pdfplumber
import pytesseract
from PIL import Image, ImageFilter, ImageEnhance
from docx import Document
import io
import cv2
import numpy as np

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# ==============================
# IMAGE PREPROCESSING
# ==============================
def preprocess_image(image: Image.Image) -> list[Image.Image]:
    """
    Returns a list of preprocessed PIL Images to try with Tesseract.
    Multiple variants improve chances of a clean OCR result.
    """
    variants = []

    # --- Variant 1: Raw upscaled image (baseline) ---
    w, h = image.size
    upscaled = image.resize((w * 2, h * 2), Image.LANCZOS)
    variants.append(upscaled.convert("L"))  # grayscale

    # --- Variant 2: OpenCV adaptive threshold ---
    try:
        # PIL is RGB — convert correctly (not BGR)
        img_array = np.array(image.convert("RGB"))
        gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)

        # Upscale
        gray = cv2.resize(gray, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)

        # Denoise
        denoised = cv2.GaussianBlur(gray, (3, 3), 0)

        # Adaptive threshold
        thresh = cv2.adaptiveThreshold(
            denoised, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            15, 10
        )

        variants.append(Image.fromarray(thresh))

    except Exception as e:
        print(f"⚠️ OpenCV preprocessing failed, skipping variant: {e}")

    # --- Variant 3: PIL sharpening + contrast boost ---
    try:
        enhanced = image.convert("L")
        enhanced = ImageEnhance.Contrast(enhanced).enhance(2.0)
        enhanced = ImageEnhance.Sharpness(enhanced).enhance(2.0)
        enhanced = enhanced.filter(ImageFilter.SHARPEN)
        w, h = enhanced.size
        enhanced = enhanced.resize((w * 2, h * 2), Image.LANCZOS)
        variants.append(enhanced)
    except Exception as e:
        print(f"⚠️ PIL enhancement failed, skipping variant: {e}")

    return variants


# ==============================
# OCR WITH MULTIPLE STRATEGIES
# ==============================
def run_ocr(image: Image.Image) -> str:
    """
    Tries multiple preprocessed variants × multiple PSM configs.
    Returns the longest (most complete) result.
    """
    psm_configs = [
        r"--oem 3 --psm 4",   # single column of text
        r"--oem 3 --psm 6",   # uniform block of text
        r"--oem 3 --psm 3",   # fully automatic (default)
        r"--oem 3 --psm 11",  # sparse text
    ]

    preprocessed_variants = preprocess_image(image)
    best_text = ""

    for variant in preprocessed_variants:
        for config in psm_configs:
            try:
                text = pytesseract.image_to_string(variant, config=config)
                cleaned = text.strip()
                if len(cleaned) > len(best_text):
                    best_text = cleaned
            except Exception as e:
                print(f"⚠️ OCR attempt failed (config={config}): {e}")

    return best_text


# ==============================
# MAIN FUNCTION
# ==============================
def extract_text(file) -> str:

    filename = file.filename.lower()
    contents = file.file.read()
    file_stream = io.BytesIO(contents)

    text = ""

    # -------- PDF --------
    if filename.endswith(".pdf"):
        with pdfplumber.open(file_stream) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

                # Fallback: if a PDF page has no text layer, OCR its image
                if not page_text or not page_text.strip():
                    print(f"⚠️ Page has no text layer — attempting image OCR fallback")
                    try:
                        for img_obj in page.images:
                            pil_img = page.to_image(resolution=200).original
                            ocr_text = run_ocr(pil_img)
                            if ocr_text:
                                text += ocr_text + "\n"
                                break
                    except Exception as e:
                        print(f"⚠️ PDF image OCR fallback failed: {e}")

        print("📄 PDF TEXT:\n", text[:500])
        return text.strip()

    # -------- IMAGE --------
    elif filename.endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".webp")):
        try:
            image = Image.open(file_stream)

            # Handle RGBA / palette mode images
            if image.mode in ("RGBA", "P", "CMYK"):
                image = image.convert("RGB")

            text = run_ocr(image)

            if not text.strip():
                print("⚠️ OCR returned empty text — check image quality or resolution")
            else:
                print("🖼️ IMAGE OCR TEXT:\n", text[:500])

            return text.strip()

        except Exception as e:
            print(f"❌ Image extraction failed: {e}")
            return ""

    # -------- DOCX --------
    elif filename.endswith(".docx"):
        try:
            doc = Document(file_stream)
            for para in doc.paragraphs:
                text += para.text + "\n"

            # Also extract text from tables (common in lab reports)
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text += row_text + "\n"

            print("📄 DOCX TEXT:\n", text[:500])
            return text.strip()

        except Exception as e:
            print(f"❌ DOCX extraction failed: {e}")
            return ""

    else:
        print(f"❌ Unsupported file type: {filename}")
        return ""